"""Document ingestion pipeline for RAG."""

import os
import uuid

from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.core.config import settings
from app.services.embedding_service import get_embeddings
from app.services.retriever_service import get_collection

_META_KEYWORDS = [
    "chatbot",
    "chunking",
    "cấu trúc dữ liệu",  # "cau truc du lieu"
]


def _is_meta_section(title: str) -> bool:
    t = title.lower()
    return any(kw in t for kw in _META_KEYWORDS)


def _load_docx_by_sections(file_path: str) -> list:
    """Load a DOCX file while preserving heading sections."""
    from docx import Document as DocxDocument
    from langchain_core.documents import Document

    doc = DocxDocument(file_path)
    sections = []
    current_title = ""
    current_lines = []

    def _flush():
        if not current_lines:
            return
        if _is_meta_section(current_title):
            return
        content = (current_title + "\n" + "\n".join(current_lines)).strip()
        sections.append(Document(
            page_content=content,
            metadata={
                "source": file_path,
                "section_title": current_title,
            },
        ))

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = (para.style.name or "").lower() if para.style else ""
        is_heading = "heading" in style_name

        if is_heading:
            _flush()
            current_lines = []
            current_title = text
        else:
            current_lines.append(text)

    _flush()

    if not sections:
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [Document(
            page_content=full_text,
            metadata={"source": file_path, "section_title": ""},
        )]

    return sections


def _load_doc_legacy(file_path: str) -> list:
    """Load a legacy .doc file via MS Word COM automation (Windows only)."""
    import sys
    from langchain_core.documents import Document

    if sys.platform != "win32":
        raise ValueError(
            f"Cannot read '{os.path.basename(file_path)}' (.doc format): "
            "MS Word COM automation is only available on Windows. "
            "Convert the file to .docx and re-ingest."
        )

    try:
        import pythoncom
        import win32com.client

        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(os.path.abspath(file_path))
            text = doc.Content.Text
            doc.Close(False)
        finally:
            word.Quit()
            pythoncom.CoUninitialize()

        return [Document(
            page_content=text,
            metadata={"source": file_path, "section_title": ""},
        )]

    except ImportError:
        raise ValueError(
            f"Cannot read '{os.path.basename(file_path)}' (.doc format). "
            "Install pywin32 (pip install pywin32) or convert the file to .docx first."
        )
    except Exception as e:
        raise ValueError(f"Failed to open .doc file via MS Word: {e}")


def load_document(file_path: str) -> list:
    """Load a supported file and return LangChain documents."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return PyPDFLoader(file_path).load()

    if ext == ".docx":
        return _load_docx_by_sections(file_path)

    if ext == ".doc":
        return _load_doc_legacy(file_path)

    if ext == ".txt":
        return TextLoader(file_path, encoding="utf-8").load()

    raise ValueError(f"Unsupported file format: {ext}")


def split_documents(documents: list) -> list:
    """Split documents into retrieval-sized chunks."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        add_start_index=True,
        separators=[
            "\n# ", "\n## ", "\n### ", "\n#### ",
            "\nChapter ", "\nSection ", "\nPart ", "\nArticle ",
            "\n- ", "\n* ", "\n• ",
            "\n\n",
            "\n",
            ". ",
            "? ",
            "! ",
            " ",
            "",
        ],
    )
    return splitter.split_documents(documents)


def ingest_file(
    file_path: str,
    title: str = None,
    category: str = "general",
    access_level: str = "all",
    department: str = "general",
) -> dict:
    """Ingest one file into the vector store."""
    documents = load_document(file_path)
    if not documents:
        return {"status": "error", "message": "Failed to read file"}

    chunks = split_documents(documents)
    if not chunks:
        return {"status": "error", "message": "Failed to create chunks"}

    file_name = os.path.basename(file_path)
    if title is None:
        title = os.path.splitext(file_name)[0]

    embeddings = get_embeddings()
    texts = [chunk.page_content for chunk in chunks]
    vectors = embeddings.embed_documents(texts)

    collection = get_collection()

    ids = []
    metadatas = []
    valid_texts = []
    valid_vectors = []

    for i, chunk in enumerate(chunks):
        chunk_id = f"{file_name}_{uuid.uuid4().hex[:8]}"
        ids.append(chunk_id)
        page = chunk.metadata.get("page", i + 1)
        section_title = chunk.metadata.get("section_title", "")

        metadatas.append({
            "title": title,
            "category": category,
            "access_level": access_level,
            "department": department,
            "source_file": file_name,
            "page": page,
            "section_title": section_title,
        })
        valid_texts.append(texts[i])
        valid_vectors.append(vectors[i])

    collection.add(
        ids=ids,
        documents=valid_texts,
        embeddings=valid_vectors,
        metadatas=metadatas,
    )

    return {
        "status": "success",
        "file": file_name,
        "title": title,
        "chunks_created": len(valid_texts),
        "category": category,
        "access_level": access_level,
    }


def ingest_directory(directory: str = None) -> list[dict]:
    """Ingest all supported files in a directory."""
    if directory is None:
        directory = settings.DOCS_DIR

    if not os.path.exists(directory):
        return [{"status": "error", "message": f"Directory {directory} does not exist"}]

    results = []
    supported_exts = {".pdf", ".docx", ".doc", ".txt"}

    for filename in os.listdir(directory):
        ext = os.path.splitext(filename)[1].lower()
        if ext in supported_exts:
            file_path = os.path.join(directory, filename)
            try:
                result = ingest_file(file_path)
                results.append(result)
            except Exception as e:
                results.append({
                    "status": "error",
                    "file": filename,
                    "message": str(e),
                })

    return results


def get_collection_stats() -> dict:
    """Return vector collection stats."""
    collection = get_collection()
    return {
        "collection_name": settings.CHROMA_COLLECTION_NAME,
        "total_chunks": collection.count(),
    }
